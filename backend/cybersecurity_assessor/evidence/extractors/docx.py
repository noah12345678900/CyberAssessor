"""DOCX extractor (python-docx, lazily imported).

Walks paragraphs and table cells in document order. We deliberately do
NOT preserve formatting — the downstream LLM only needs the words to
cite the document, and stripping styles keeps the cache key small and
the text snapshot diffable in tests.
"""

from __future__ import annotations

from pathlib import PurePosixPath
from typing import BinaryIO

from ...models import EvidenceKind
from .base import ExtractedDoc, ExtractorError, register, resolve_doc_number


@register(".docx")
def extract_docx(stream: BinaryIO, name: str) -> ExtractedDoc:
    """Extract paragraph + table text from a Word document."""
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ExtractorError(
            "python-docx is not installed — add it to backend/pyproject.toml "
            "to extract DOCX evidence."
        ) from exc

    stem = PurePosixPath(name).stem

    try:
        doc = Document(stream)
    except Exception as exc:  # pragma: no cover - python-docx errors vary
        raise ExtractorError(f"python-docx failed on {name}: {exc}") from exc

    chunks: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            chunks.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    text = "\n".join(chunks)
    title = (doc.core_properties.title or "").strip() or stem

    return ExtractedDoc(
        text=text,
        title=title,
        doc_number=resolve_doc_number(name, title, text),
        kind=EvidenceKind.DOCX,
    )
